"""최종보고서.docx 생성 스크립트"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 기본 여백 설정 ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── 색상 상수 ──────────────────────────────────────────────
PURPLE     = RGBColor(0x6c, 0x63, 0xff)
DARK       = RGBColor(0x1a, 0x1a, 0x2e)
GRAY       = RGBColor(0x6b, 0x72, 0x80)
WHITE      = RGBColor(0xff, 0xff, 0xff)
LIGHT_GRAY = RGBColor(0x37, 0x41, 0x51)


# ── 헬퍼: 단락 스타일 적용 ──────────────────────────────────
def set_para_fmt(para, space_before=0, space_after=6, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)


def heading1(text, num=None):
    """섹션 번호 + 제목"""
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=18, space_after=8)
    if num:
        run = p.add_run(f"{num}. ")
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = PURPLE
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = DARK
    # 하단 테두리
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '6c63ff')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(text):
    """◆ 소제목"""
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=12, space_after=4)
    run = p.add_run("◆ " + text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    return p


def heading3(text):
    """소소제목"""
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=8, space_after=3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = LIGHT_GRAY
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    set_para_fmt(p, space_after=4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = LIGHT_GRAY
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_para_fmt(p, space_after=3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = LIGHT_GRAY
    return p


def number_item(text, num):
    p = doc.add_paragraph()
    set_para_fmt(p, space_after=3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{num}. ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = PURPLE
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = LIGHT_GRAY
    return p


def code_line(text):
    p = doc.add_paragraph()
    set_para_fmt(p, space_after=2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size   = Pt(9.5)
    run.font.name   = 'Consolas'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 배경색 적용
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    rPr.append(shd)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 배경색 보라
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '6c63ff')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'F9FAFB' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.color.rgb = LIGHT_GRAY

    # 열 너비
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 표 아래 여백
    return table


def page_break():
    doc.add_page_break()


def divider():
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=4, space_after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bottom)
    pPr.append(pBdr)


def badge_line(items):
    """뱃지처럼 보이는 쉼표 구분 기술 스택"""
    p = doc.add_paragraph()
    set_para_fmt(p, space_after=4)
    p.paragraph_format.left_indent = Cm(0.5)
    for i, item in enumerate(items):
        if i:
            sep = p.add_run("  ·  ")
            sep.font.size = Pt(10)
            sep.font.color.rgb = GRAY
        r = p.add_run(item)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = PURPLE
    return p


def info_box(title_text, sub_text=None):
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=4, space_after=8)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(title_text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = PURPLE
    if sub_text:
        p2 = doc.add_paragraph()
        set_para_fmt(p2, space_after=8)
        p2.paragraph_format.left_indent = Cm(0.5)
        r2 = p2.add_run(sub_text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = GRAY


def prompt_box(label, text):
    p0 = doc.add_paragraph()
    set_para_fmt(p0, space_before=4, space_after=2)
    p0.paragraph_format.left_indent = Cm(0.5)
    r0 = p0.add_run(label)
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.color.rgb = PURPLE

    p = doc.add_paragraph()
    set_para_fmt(p, space_after=8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Malgun Gothic'
    run.font.color.rgb = DARK
    # 배경 박스
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0EEFF')
    rPr.append(shd)


# ══════════════════════════════════════════════════════════
#  표지
# ══════════════════════════════════════════════════════════
p_tag = doc.add_paragraph()
p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_fmt(p_tag, space_before=80, space_after=20)
r = p_tag.add_run("TERM PROJECT FINAL REPORT")
r.font.size = Pt(10)
r.font.color.rgb = PURPLE
r.bold = True

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_fmt(p_title, space_before=0, space_after=8)
r = p_title.add_run("실시간 경제 소식 대시보드")
r.font.size = Pt(26)
r.bold = True
r.font.color.rgb = DARK

p_en = doc.add_paragraph()
p_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_fmt(p_en, space_after=50)
r = p_en.add_run("Real-time Economic News Dashboard")
r.font.size = Pt(13)
r.font.color.rgb = GRAY

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_fmt(p_info, space_after=6)
r = p_info.add_run("학번   20251429")
r.font.size = Pt(13)
r.font.color.rgb = DARK

p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_fmt(p_name, space_after=6)
r = p_name.add_run("이름   박동우")
r.font.size = Pt(13)
r.font.color.rgb = DARK
r.bold = True

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_fmt(p_date, space_after=0)
r = p_date.add_run("2026. 05. 25.")
r.font.size = Pt(11)
r.font.color.rgb = GRAY

page_break()


# ══════════════════════════════════════════════════════════
#  SECTION 01 — 텀프로젝트 제목
# ══════════════════════════════════════════════════════════
heading1("텀프로젝트 제목", num="1")

heading2("프로젝트명")
info_box("실시간 경제 소식 대시보드",
         "Real-time Economic News Dashboard with AI Summarization")

heading2("프로젝트 개요")
body("다수의 경제 뉴스 RSS 피드와 금융 데이터 API를 통합 수집하고, "
     "대형 언어 모델(LLM)로 뉴스를 자동 요약하여 Streamlit 기반 단일 대시보드에 "
     "표시하는 실시간 경제 정보 플랫폼입니다.")

heading2("사용 기술 스택")
badge_line(["Python 3.11", "Streamlit", "SQLite", "APScheduler",
            "OpenRouter API", "yfinance", "feedparser", "Plotly", "openai SDK"])

divider()


# ══════════════════════════════════════════════════════════
#  SECTION 02 — 목적과 필요성
# ══════════════════════════════════════════════════════════
heading1("텀프로젝트 목적과 필요성", num="2")

heading2("개발 배경 및 문제 인식")
body("현대 경제 환경에서 개인 투자자와 직장인은 수많은 경제 뉴스와 시장 데이터를 "
     "매일 소화해야 합니다. 그러나 다음과 같은 구조적 불편함이 존재합니다.")
bullet("정보 분산 문제 — 연합뉴스, 한국경제, Reuters, CNN 등 주요 경제 뉴스가 각각의 "
       "사이트에 흩어져 있어 일일이 방문해야 함")
bullet("정보 과부하 — 하루에도 수백 건의 뉴스가 생성되어 핵심 내용만 빠르게 파악하기 어려움")
bullet("시장 지표 분산 — KOSPI·NASDAQ·환율·원자재 등 주요 지표를 한 화면에서 비교하기 어려움")
bullet("시장 심리 파악의 어려움 — CNN Fear & Greed Index 같은 심리 지표를 별도로 탐색해야 함")

heading2("프로젝트 목적")
number_item("여러 경제 뉴스 소스를 자동 수집하여 한 화면에서 제공", 1)
number_item("AI(LLM)를 활용해 뉴스 제목을 한국어 2~3문장으로 자동 요약", 2)
number_item("주요 주가 지수·환율·원자재 등 시장 지표 8종을 실시간 시각화", 3)
number_item("10분 주기 자동 갱신으로 항상 최신 정보 유지", 4)
number_item("정적 HTML 내보내기를 통해 서버 없이도 공유 가능한 보고서 생성", 5)

heading2("기대 효과 (KPI)")
add_table(
    headers=["지표", "목표값"],
    rows=[
        ["페이지 초기 로드 시간",      "≤ 3초"],
        ["시장 데이터 갱신 주기",      "≤ 10분"],
        ["AI 뉴스 요약 생성 시간",     "≤ 30초 / 건"],
        ["뉴스 수집 소스",            "4개 이상 RSS 피드"],
        ["시장 지표 커버리지",
         "KOSPI · KOSDAQ · S&P 500 · NASDAQ · Dow Jones · USD/KRW · Gold · WTI Oil"],
    ],
    col_widths=[6, 10]
)

divider()


# ══════════════════════════════════════════════════════════
#  SECTION 03 — 개발 내용
# ══════════════════════════════════════════════════════════
heading1("텀프로젝트 개발내용", num="3")

heading2("시스템 아키텍처")
body("전체 시스템은 수집 → 저장 → 표시의 단방향 파이프라인 구조로 설계되었습니다.")

# 흐름 표
add_table(
    headers=["단계", "구성요소", "역할"],
    rows=[
        ["TRIGGER",  "Scheduler (APScheduler)",           "10분 주기 작업 실행"],
        ["COLLECT",  "Modules (news / market / summarizer)", "데이터 수집 및 AI 요약"],
        ["STORE",    "SQLite DB (economy.db)",             "뉴스·시장 데이터 영속 저장"],
        ["DISPLAY",  "Dashboard (Streamlit)",              "읽기 전용 시각화"],
    ],
    col_widths=[3, 7, 6]
)

heading2("디렉토리 구조")
for line in [
    "project/",
    "  modules/",
    "    news.py          # RSS 피드 수집 (feedparser)",
    "    market.py        # yfinance 시장 데이터 + CNN Fear & Greed",
    "    summarizer.py    # OpenRouter LLM 뉴스 요약",
    "  db/",
    "    schema.py        # SQLite 테이블 초기화",
    "    helpers.py       # CRUD 헬퍼 함수 7종",
    "  scheduler/",
    "    jobs.py          # APScheduler 10분 주기 작업",
    "  dashboard/",
    "    app.py           # Streamlit 메인 앱",
    "    export_html.py   # 정적 index.html 생성",
    "  tests/             # pytest 테스트 스위트",
]:
    code_line(line)

heading2("주요 모듈 개발 내용")

heading3("① 뉴스 수집 모듈 (modules/news.py)")
body("feedparser 라이브러리를 사용해 4개의 RSS 피드에서 최신 뉴스를 수집합니다. "
     "피드별로 최대 10건을 가져오며, 각 피드 실패는 try/except로 독립 처리합니다.")
add_table(
    headers=["소스명", "RSS URL"],
    rows=[
        ["연합뉴스 경제", "yna.co.kr/rss/economy.xml"],
        ["한국경제",      "rss.hankyung.com/economy.xml"],
        ["Reuters Business", "feeds.reuters.com/reuters/businessNews"],
        ["CNN Money",    "rss.cnn.com/rss/money_topstories.rss"],
    ],
    col_widths=[5, 11]
)

heading3("② 시장 데이터 모듈 (modules/market.py)")
body("yfinance의 Ticker.fast_info를 사용해 8종 종목의 현재가·전일가를 조회하고 "
     "등락률(%)을 계산합니다. CNN Fear & Greed Index는 공개 API에 브라우저 헤더를 "
     "위장하여 요청합니다.")

heading3("③ AI 요약 모듈 (modules/summarizer.py)")
body("OpenRouter API를 openai SDK의 base_url 오버라이드 방식으로 호출합니다. "
     "뉴스 제목만 입력으로 받아 LLM이 2~3문장 한국어 요약을 생성합니다.")

heading3("④ 스케줄러 (scheduler/jobs.py)")
body("APScheduler BackgroundScheduler를 사용해 뉴스·시장 데이터 수집을 10분마다 반복합니다. "
     "Streamlit의 스크립트 재실행 특성으로 인한 중복 기동을 "
     "st.session_state.initialized 플래그로 방지합니다.")

heading3("⑤ 데이터베이스 (db/)")
body("SQLite 단일 파일 DB(economy.db)에 두 테이블을 운용합니다.")
add_table(
    headers=["테이블", "주요 컬럼", "용도"],
    rows=[
        ["news",        "title, url(UNIQUE), source, summary, fetched_at", "뉴스 원문 + AI 요약 저장"],
        ["market_data", "symbol, name, price, change_pct, volume, fetched_at", "시장 지표 이력 저장"],
    ],
    col_widths=[3, 8, 5]
)

heading3("⑥ 대시보드 (dashboard/app.py)")
body("Streamlit으로 구현한 메인 화면은 다음 4개 영역으로 구성됩니다.")
bullet("사이드바 — 새로고침 버튼, 뉴스 표시 수 슬라이더(5~50), AI 요약 표시 토글")
bullet("시장 지표 섹션 — metric 카드 + Plotly 등락률 바 차트")
bullet("Fear & Greed 섹션 — Plotly 게이지 차트 + 전일/1주/1달 비교")
bullet("뉴스 섹션 — st.expander 목록 + AI 요약 인라인 표시")

divider()


# ══════════════════════════════════════════════════════════
#  SECTION 04 — 이용대상 및 업무 분석
# ══════════════════════════════════════════════════════════
heading1("이 앱의 이용대상 및 업무 분석", num="4")

heading2("주요 이용 대상")
add_table(
    headers=["페르소나", "특징", "핵심 니즈"],
    rows=[
        ["개인 투자자",   "30~50대, 주식·ETF·원자재 보유자",       "실시간 등락률 + 한국어 뉴스 요약"],
        ["직장인",        "출퇴근 중 경제 뉴스를 빠르게 파악",      "짧은 AI 요약, 핵심 정보 집약"],
        ["금융 전공 학생", "리서치·과제·스터디 목적의 학부·대학원생", "차트 시각화, 데이터 내보내기"],
    ],
    col_widths=[4, 7, 5.5]
)

heading2("업무 분석 (Use Case)")

heading3("UC-01. 경제 뉴스 조회")
bullet("행위자: 모든 사용자")
bullet("흐름: 대시보드 접속 → 뉴스 목록 확인 → 제목 클릭 → AI 요약 확인 → 원문 링크 이동")
bullet("전제조건: 스케줄러가 1회 이상 수집 완료")

heading3("UC-02. 시장 지표 모니터링")
bullet("행위자: 개인 투자자")
bullet("흐름: 대시보드 접속 → 시장 지표 카드 확인 → 바 차트로 등락률 비교 → Fear & Greed 게이지 확인")
bullet("전제조건: yfinance API 정상 응답")

heading3("UC-03. 데이터 내보내기 (공유)")
bullet("행위자: 학생, 직장인")
bullet("흐름: export_html.py 실행 → index.html 생성 → 이메일·메신저 등으로 파일 공유")
bullet("특이사항: 수신자는 브라우저만 있으면 별도 설치 없이 열람 가능")

heading3("UC-04. 자동 갱신 유지")
bullet("행위자: 시스템 (APScheduler)")
bullet("흐름: 10분 경과 → 뉴스·시장 데이터 자동 수집 → DB 갱신 → 다음 새로고침 시 반영")
bullet("특이사항: 사용자 개입 없이 항상 최신 상태 유지")

heading2("시스템 경계 및 제외 기능")
body("MVP 범위를 초과하거나 현실적 제약으로 이번 버전에서 제외한 기능은 다음과 같습니다.")
add_table(
    headers=["제외 기능", "제외 이유"],
    rows=[
        ["뉴스 본문 크롤링",   "사이트별 HTML 구조 상이, 유지 비용 높음 → 제목 기반 요약으로 대체"],
        ["감성 분석 (긍/부정)", "추가 프롬프트 설계 및 API 비용 증가"],
        ["키워드 필터·검색",   "DB 쿼리 확장 및 UI 설계 범위 초과"],
        ["포트폴리오 손익 추적", "사용자 계정 체계 및 인증 구현 필요"],
        ["이메일 알림",        "외부 SMTP 서버 연동 범위 초과"],
        ["모바일 앱",         "Streamlit 웹 기반으로 반응형 지원, 별도 앱 빌드 불필요"],
    ],
    col_widths=[4.5, 12]
)

divider()


# ══════════════════════════════════════════════════════════
#  SECTION 05 — 사용한 AI 프롬프트
# ══════════════════════════════════════════════════════════
heading1("사용한 AI 프롬프트", num="5")

heading2("프롬프트 위치 및 역할")
body("AI 프롬프트는 project/modules/summarizer.py의 summarize() 함수에 정의되어 있습니다. "
     "뉴스 제목을 입력받아 한국어 요약문을 반환하는 단일 프롬프트입니다.")

heading2("마인드맵 구조 (AI 프롬프트 설계 요소)")
body("아래 6개 항목이 원형으로 배열된 마인드맵으로, 각 항목은 중앙 'AI 프롬프트' 노드와 연결됩니다.")
add_table(
    headers=["카테고리", "세부 항목"],
    rows=[
        ["프롬프트 내용", "한국어 출력 고정 / 2~3문장 요약 / 입력: {title} / 간결한 톤"],
        ["LLM 설정",     "gpt-4o-mini / OpenRouter API / max_tokens 200 / temperature 0.3"],
        ["호출 흐름",    "APScheduler / 10분 주기 / 미요약 필터 / DB 저장"],
        ["비용 제어",    "회당 5건 상한 / 저가 모델 / 토큰 상한 200"],
        ["출력 표시",    "st.info() 박스 / expander UI / index.html"],
        ["설계 의도",    "제목만 입력 / 단일 user 메시지 / 크롤링 불필요"],
    ],
    col_widths=[4, 12.5]
)

heading2("실제 사용 프롬프트")
prompt_box(
    "PROMPT — summarizer.py : summarize()",
    "다음 경제 뉴스 제목을 읽고, 핵심 내용을 한국어로 2~3문장으로 간결하게 요약하세요.\n\n"
    "제목: {title}"
)

heading2("프롬프트 설계 의도")
bullet("언어 고정 — '한국어로'를 명시하여 영문 제목 입력 시에도 한국어 요약 생성")
bullet("길이 지정 — '2~3문장'으로 출력 분량을 명시하여 과도한 설명 방지")
bullet("톤 지정 — '간결하게'로 불필요한 수식어 없이 핵심만 서술하도록 유도")
bullet("단순 구조 — 시스템 프롬프트 없이 단일 user 메시지만 사용, API 비용 최소화")
bullet("제목 기반 — 본문 크롤링 없이 제목만으로 요약 생성, 사이트 차단·구조 변경에 강건")

heading2("LLM 호출 파라미터")
add_table(
    headers=["파라미터", "값", "역할"],
    rows=[
        ["model",       "openai/gpt-4o-mini",  "OpenRouter 경유 · 저가 고성능 모델 (기본값, 변경 가능)"],
        ["max_tokens",  "200",                  "짧은 요약 강제 · API 과금 상한 역할"],
        ["temperature", "0.3",                  "낮은 값 → 일관되고 사실적인 출력 보장"],
        ["role",        "user",                 "단일 user 메시지 구조 (system 프롬프트 없음)"],
    ],
    col_widths=[3.5, 4, 9]
)

heading2("전체 호출 코드")
for line in [
    "def summarize(title: str, model: str = 'openai/gpt-4o-mini') -> str | None:",
    "    prompt = (",
    '        "다음 경제 뉴스 제목을 읽고, 핵심 내용을 한국어로 2~3문장으로 간결하게 요약하세요.\\n\\n"',
    '        f"제목: {title}"',
    "    )",
    "    resp = _get_client().chat.completions.create(",
    "        model=model,",
    '        messages=[{"role": "user", "content": prompt}],',
    "        max_tokens=200,",
    "        temperature=0.3,",
    "    )",
    "    return resp.choices[0].message.content.strip()",
]:
    code_line(line)

heading2("비용 제어 전략")
bullet("회당 최대 5건 제한 — get_unsummarized_news(limit=5)로 한 사이클당 LLM 호출 횟수 상한")
bullet("저가 모델 기본값 — gpt-4o-mini는 고성능 대비 비용이 낮아 대량 처리에 적합")
bullet("max_tokens 200 — 출력 토큰을 제한해 예상치 못한 과금 방지")
bullet("실패 시 재시도 방지 — 오류 발생 건은 summary = None 유지, 다음 회차에 자동 재시도")


# ══════════════════════════════════════════════════════════
#  저장
# ══════════════════════════════════════════════════════════
out_path = r"C:\재출\최종보고서.docx"
doc.save(out_path)
print(f"저장 완료: {out_path}")
